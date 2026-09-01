import json
import pathlib
import re
import sys

from docx import Document


def clean(value):
    return re.sub(r"[ \t]+", " ", value.replace("\xa0", " ")).strip()


root = pathlib.Path(sys.argv[1])
conversion = json.loads((root / "conversion-results.json").read_text(encoding="utf-8-sig"))
documents = []
for entry in conversion:
    if entry.get("status") != "converted" or not entry.get("output"):
        documents.append({"source": entry["source"], "status": entry.get("status"), "error": entry.get("error")})
        continue
    path = pathlib.Path(entry["output"])
    try:
        document = Document(path)
        paragraphs = [clean(paragraph.text) for paragraph in document.paragraphs if clean(paragraph.text)]
        tables = []
        for table in document.tables:
            rows = []
            for row in table.rows:
                rows.append([clean(cell.text) for cell in row.cells])
            tables.append(rows)
        documents.append({
            "source": entry["source"],
            "status": "extracted",
            "paragraphs": paragraphs,
            "tables": tables,
            "table_shapes": [[len(table), max((len(row) for row in table), default=0)] for table in tables],
        })
    except Exception as error:
        documents.append({"source": entry["source"], "status": "error", "error": str(error)})

payload = json.dumps(documents, ensure_ascii=False)
if len(sys.argv) > 2:
    pathlib.Path(sys.argv[2]).write_text(payload, encoding="utf-8")
else:
    print(payload)
